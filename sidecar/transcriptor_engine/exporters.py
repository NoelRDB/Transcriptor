from __future__ import annotations

import csv
import io
import json
from html import escape
from pathlib import Path
from typing import Any

from .unicode_text import repair_data, sanitize_text


def subtitle_time(milliseconds: int, separator: str) -> str:
    value = max(0, round(milliseconds))
    hours, remainder = divmod(value, 3_600_000)
    minutes, remainder = divmod(remainder, 60_000)
    seconds, millis = divmod(remainder, 1000)
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}{separator}{millis:03d}"


def render(project: dict[str, Any], export_format: str) -> str:
    # Repair text saved by early builds after UTF-8 was decoded as a Western
    # single-byte encoding. Sanitizing surrogates alone would keep visible
    # sequences such as "mÃ¡s" in the exported document.
    project = repair_data(project)
    segments = sorted(
        project.get("segments", []), key=lambda segment: (segment["startMs"], segment.get("order", 0))
    )
    settings = project.get("settings", {})
    subtitle_segments = _subtitle_cues(
        segments,
        int(settings.get("subtitleLineLength", 42)),
        int(settings.get("subtitleMaxLines", 2)),
    )
    if export_format == "txt":
        return "\n\n".join(
            f"{segment.get('speaker') + ': ' if segment.get('speaker') else ''}{segment['text'].strip()}"
            for segment in segments
        )
    if export_format == "srt":
        blocks = []
        for index, segment in enumerate(subtitle_segments, 1):
            start = max(0, int(segment["startMs"]))
            end = max(start + 1, int(segment["endMs"]))
            speaker = f"[{segment['speaker']}] " if segment.get("speaker") else ""
            blocks.append(
                f"{index}\n{subtitle_time(start, ',')} --> {subtitle_time(end, ',')}\n"
                f"{speaker}{segment['text'].strip()}"
            )
        return "\n\n".join(blocks) + "\n"
    if export_format == "vtt":
        blocks = []
        for segment in subtitle_segments:
            start = max(0, int(segment["startMs"]))
            end = max(start + 1, int(segment["endMs"]))
            blocks.append(
                f"{subtitle_time(start, '.')} --> {subtitle_time(end, '.')}\n{segment['text'].strip()}"
            )
        return "WEBVTT\n\n" + "\n\n".join(blocks) + "\n"
    if export_format == "json":
        clean = {key: value for key, value in project.items() if key != "mediaUrl"}
        return json.dumps({"version": 1, "project": clean}, ensure_ascii=False, indent=2)
    if export_format == "csv":
        output = io.StringIO(newline="")
        writer = csv.writer(output)
        writer.writerow(["start_ms", "end_ms", "speaker", "text"])
        writer.writerows(
            [
                [segment["startMs"], segment["endMs"], segment.get("speaker", ""), segment["text"]]
                for segment in segments
            ]
        )
        return output.getvalue()
    raise ValueError(f"Formato de exportación no compatible: {export_format}")


def _subtitle_cues(
    segments: list[dict[str, Any]], line_length: int = 42, max_lines: int = 2
) -> list[dict[str, Any]]:
    line_length = min(80, max(18, line_length))
    max_lines = min(3, max(1, max_lines))
    cues: list[dict[str, Any]] = []
    for segment in segments:
        words = str(segment.get("text") or "").split()
        if not words:
            continue
        groups: list[list[str]] = []
        current_group: list[str] = []
        current_line = ""
        lines = 1
        for word in words:
            candidate = f"{current_line} {word}".strip()
            if current_line and len(candidate) > line_length:
                if lines >= max_lines:
                    groups.append(current_group)
                    current_group = []
                    lines = 1
                else:
                    lines += 1
                current_line = word
            else:
                current_line = candidate
            current_group.append(word)
        if current_group:
            groups.append(current_group)
        total_weight = sum(max(1, len(" ".join(group))) for group in groups)
        start_ms = max(0, int(segment.get("startMs") or 0))
        end_ms = max(start_ms + 1, int(segment.get("endMs") or start_ms + 1))
        elapsed_weight = 0
        for index, group in enumerate(groups):
            cue_start = round(start_ms + (end_ms - start_ms) * elapsed_weight / total_weight)
            elapsed_weight += max(1, len(" ".join(group)))
            cue_end = (
                end_ms
                if index == len(groups) - 1
                else round(start_ms + (end_ms - start_ms) * elapsed_weight / total_weight)
            )
            lines_output: list[str] = []
            line = ""
            for word in group:
                candidate = f"{line} {word}".strip()
                if line and len(candidate) > line_length:
                    lines_output.append(line)
                    line = word
                else:
                    line = candidate
            if line:
                lines_output.append(line)
            cues.append(
                {
                    **segment,
                    "startMs": cue_start,
                    "endMs": max(cue_start + 1, cue_end),
                    "text": "\n".join(lines_output[:max_lines]),
                }
            )
    return cues


def export_to(project: dict[str, Any], export_format: str, output_path: str) -> None:
    target = Path(output_path).resolve()
    if target.suffix.lower() != f".{export_format}":
        target = target.with_suffix(f".{export_format}")
    target.parent.mkdir(parents=True, exist_ok=True)
    project = repair_data(project)
    if export_format == "docx":
        _export_docx(project, target)
        return
    if export_format == "pdf":
        _export_pdf(project, target)
        return
    contents = sanitize_text(render(project, export_format))
    # Some Windows editors still guess ANSI for extension-based text files.
    # A BOM makes accents and punctuation unambiguous. JSON stays plain UTF-8
    # for maximum parser interoperability.
    encoding = "utf-8" if export_format == "json" else "utf-8-sig"
    target.write_text(contents, encoding=encoding, errors="strict", newline="")


def _export_docx(project: dict[str, Any], target: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    document = Document()
    document.core_properties.title = sanitize_text(str(project.get("name") or "Transcripción"))
    document.add_heading(document.core_properties.title, level=0)
    metadata = document.add_paragraph()
    metadata.add_run("Duración: ").bold = True
    metadata.add_run(subtitle_time(int(project.get("durationMs") or 0), ".")[:-4])
    metadata.add_run("    Idioma: ").bold = True
    metadata.add_run(sanitize_text(str(project.get("detectedLanguage") or project.get("language") or "—")))
    for segment in _sorted_segments(project):
        paragraph = document.add_paragraph()
        timestamp = paragraph.add_run(f"[{subtitle_time(int(segment['startMs']), '.')[:-4]}] ")
        timestamp.bold = True
        timestamp.font.size = Pt(9)
        if segment.get("speaker"):
            speaker = paragraph.add_run(f"{sanitize_text(str(segment['speaker']))}: ")
            speaker.bold = True
        paragraph.add_run(sanitize_text(str(segment.get("text") or "").strip()))
    document.save(target)


def _export_pdf(project: dict[str, Any], target: Path) -> None:
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "TranscriptBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        alignment=TA_LEFT,
        spaceAfter=4,
    )
    meta = ParagraphStyle("TranscriptMeta", parent=body, fontSize=8, textColor="#667074")
    document = SimpleDocTemplate(
        str(target),
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title=sanitize_text(str(project.get("name") or "Transcripción")),
    )
    story: list[Any] = [
        Paragraph(escape(sanitize_text(str(project.get("name") or "Transcripción"))), styles["Title"]),
        Paragraph(
            f"Duración: {subtitle_time(int(project.get('durationMs') or 0), '.')[:-4]} · "
            "Idioma: "
            + escape(
                sanitize_text(str(project.get("detectedLanguage") or project.get("language") or "—"))
            ),
            meta,
        ),
        Spacer(1, 5 * mm),
    ]
    for segment in _sorted_segments(project):
        timestamp = subtitle_time(int(segment["startMs"]), ".")[:-4]
        speaker = (
            f"<b>{escape(sanitize_text(str(segment['speaker'])))}</b>: "
            if segment.get("speaker")
            else ""
        )
        text = escape(sanitize_text(str(segment.get("text") or "").strip())).replace("\n", "<br/>")
        story.append(Paragraph(f"<font color='#6b7779'>[{timestamp}]</font> {speaker}{text}", body))
    document.build(story)


def _sorted_segments(project: dict[str, Any]) -> list[dict[str, Any]]:
    return sorted(
        project.get("segments", []),
        key=lambda segment: (int(segment.get("startMs") or 0), int(segment.get("order") or 0)),
    )
